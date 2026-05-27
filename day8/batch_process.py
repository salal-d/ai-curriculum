import os
import anthropic
from dotenv import load_dotenv
from docx import Document
from datetime import datetime
import json

load_dotenv()
client = anthropic.Anthropic()

def extract_docx(filepath):
    doc = Document(filepath)
    return "\n".join(p.text for p in doc.paragraphs if p.text.strip())

def summarize(text, filename):
    response = client.messages.create(
        model="claude-sonnet-4-5",
        max_tokens=512,
        messages=[{
            "role": "user",
            "content": f"Summarize this document in 3 bullet points, then list any action items:\n\n{text}"
        }]
    )
    return response.content[0].text

# Create a second sample doc to batch process
doc2 = Document()
doc2.add_heading("Client Onboarding Checklist", 0)
doc2.add_paragraph("Review client data sources and access requirements.")
doc2.add_paragraph("Define project scope and success metrics.")
doc2.add_paragraph("Schedule kickoff meeting with stakeholders.")
doc2.add_paragraph("Set up secure data transfer protocols.")
doc2.save("onboarding_checklist.docx")

# Batch process all docx files in current directory
results = []
for filename in os.listdir("."):
    if filename.endswith(".docx"):
        print(f"Processing {filename}...")
        text = extract_docx(filename)
        summary = summarize(text, filename)
        results.append({
            "file": filename,
            "summary": summary
        })
        print(f"Done: {filename}\n")

# Save results
timestamp = datetime.now().strftime("%Y-%m-%d_%H-%M-%S")
outfile = f"batch_summary_{timestamp}.json"
with open(outfile, "w") as f:
    json.dump(results, f, indent=2)

print(f"Batch complete. {len(results)} files processed.")
print(f"Saved to {outfile}")