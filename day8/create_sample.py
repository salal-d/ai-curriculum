from docx import Document

doc = Document()
doc.add_heading("Q1 2026 AI Consulting Report", 0)

doc.add_heading("Executive Summary", 1)
doc.add_paragraph("This report summarizes AI consulting engagements for Q1 2026. Total revenue reached $127,500 across 12 projects spanning healthcare, finance, and legal sectors.")

doc.add_heading("Key Findings", 1)
doc.add_paragraph("Healthcare remains our strongest vertical with $35,000 in revenue.")
doc.add_paragraph("RAG System implementations command the highest average project value at $20,000.")
doc.add_paragraph("Document Processing projects show a 67% dissatisfaction rate and require immediate process review.")

doc.add_heading("Recommendations", 1)
doc.add_paragraph("1. Double down on Healthcare and Finance verticals.")
doc.add_paragraph("2. Audit Document Processing delivery process within 2 weeks.")
doc.add_paragraph("3. Exit Food, HR, and Energy sectors to reduce operational drag.")

doc.add_heading("Next Steps", 1)
doc.add_paragraph("Schedule post-mortems with ClearPath and Dynamo Inc by end of month.")

doc.save("sample_report.docx")
print("Created sample_report.docx")